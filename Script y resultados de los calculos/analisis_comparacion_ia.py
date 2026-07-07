import pandas as pd
import numpy as np
import os
import re
from scipy import stats
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
from docx import Document
import warnings
warnings.filterwarnings('ignore')

class ModelComparison:
    def __init__(self, data_path="casos/"):
        """
        Inicializa el analizador de comparación de modelos
        
        Args:
            data_path: Ruta donde están los archivos de casos
        """
        self.data_path = data_path
        self.models = ['Claude Opus 4.1', 'GPT-5', 'Gemini 2.5 Pro', 'DeepSeek-V3']
        self.data = []
        self.df_consolidated = None
        # Configurar matplotlib para gráficos separados
        plt.rcParams['figure.max_open_warning'] = 40
        
    def extract_text_from_docx(self, filepath):
        """
        Extrae todo el texto de un archivo .docx
        
        Args:
            filepath: Ruta del archivo .docx
        Returns:
            str: Texto completo del documento
        """
        doc = Document(filepath)
        full_text = []
        
        # Extraer texto de párrafos
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                full_text.append('\t'.join(row_data))
        
        return '\n'.join(full_text)
    
    def parse_case_file(self, filepath):
        """
        Parsea un archivo de caso individual
        
        Args:
            filepath: Ruta del archivo a parsear
        Returns:
            dict: Diccionario con los datos del caso
        """
        case_data = {}
        
        try:
            # Extraer el ID del caso del nombre del archivo
            filename = os.path.basename(filepath)
            # Patrón para CASE_CODE_C#_XXX o CASE_CODE_JAVA_XXX
            match = re.search(r'CASE_CODE_(C#|JAVA)_(\d+)', filename, re.IGNORECASE)
            if match:
                language = match.group(1).upper()
                case_number = match.group(2)
                case_data['id_caso'] = f'CASE_CODE_{language}_{case_number}'
                case_data['lenguaje'] = 'C#' if language == 'C#' else 'Java'
            
            # Leer archivo según su extensión
            if filepath.endswith('.docx'):
                content = self.extract_text_from_docx(filepath)
            else:
                with open(filepath, 'r', encoding='utf-8') as file:
                    content = file.read()
            
            # Limpiar el contenido
            content = re.sub(r'\*\*', '', content)  # Remover asteriscos de formato
            
            # Si no se extrajo del nombre del archivo, buscar en el contenido
            if 'id_caso' not in case_data:
                case_data['id_caso'] = self.extract_value(content, r'ID Caso\s+(CASE_\w+_\d+)')
            
            case_data['fecha'] = self.extract_value(content, r'Fecha\s+([\d/]+)')
            case_data['categoria_similitud'] = self.extract_value(content, r'Categoría de Similitud\s+(\w+)')
            
            # Si no se extrajo del nombre, buscar lenguaje en el contenido
            if 'lenguaje' not in case_data:
                lenguaje = self.extract_value(content, r'Lenguaje de Programación\s+(\w+#?)')
                if lenguaje:
                    # Normalizar C# que puede aparecer como C o C#
                    case_data['lenguaje'] = 'C#' if 'C' in lenguaje.upper() and 'JAVA' not in lenguaje.upper() else lenguaje
            
            # Métricas técnicas - buscar con patrón más flexible
            lines = content.split('\n')
            for i, line in enumerate(lines):
                if 'Número de líneas de código' in line:
                    numbers = re.findall(r'\d+', line)
                    if len(numbers) >= 2:
                        case_data['lineas_codigo_a'] = float(numbers[0])
                        case_data['lineas_codigo_b'] = float(numbers[1])
                    elif i + 1 < len(lines):
                        next_line = lines[i + 1]
                        numbers = re.findall(r'\d+', next_line)
                        if len(numbers) >= 2:
                            case_data['lineas_codigo_a'] = float(numbers[0])
                            case_data['lineas_codigo_b'] = float(numbers[1])
                
                if 'Número de funciones/métodos' in line:
                    numbers = re.findall(r'\d+', line)
                    if len(numbers) >= 2:
                        case_data['num_funciones_a'] = float(numbers[0])
                        case_data['num_funciones_b'] = float(numbers[1])
                    elif i + 1 < len(lines):
                        next_line = lines[i + 1]
                        numbers = re.findall(r'\d+', next_line)
                        if len(numbers) >= 2:
                            case_data['num_funciones_a'] = float(numbers[0])
                            case_data['num_funciones_b'] = float(numbers[1])
            
            # Buscar datos de cada modelo
            for model in self.models:
                model_key = model.replace(' ', '_').replace('.', '_')
                
                # Buscar valores en tablas
                case_data[f'{model_key}_tiempo'] = self.extract_table_value(content, 'Tiempo Total', model)
                case_data[f'{model_key}_clasificacion'] = self.extract_table_value(content, 'Clasificación Modelo', model)
                case_data[f'{model_key}_coincidencia'] = self.extract_table_value(content, 'Coincidencia', model)
                
                # Matriz de confusión
                case_data[f'{model_key}_vp'] = self.extract_table_value(content, 'Verdadero Positivo', model)
                case_data[f'{model_key}_vn'] = self.extract_table_value(content, 'Verdadero Negativo', model)
                case_data[f'{model_key}_fp'] = self.extract_table_value(content, 'Falso Positivo', model)
                case_data[f'{model_key}_fn'] = self.extract_table_value(content, 'Falso Negativo', model)
                
                # Corregir VN que siempre debería ser 0 en detección de similitud
                if case_data.get(f'{model_key}_vn') and case_data[f'{model_key}_vn'] > 0:
                    print(f"  ⚠️ Corrigiendo VN>0 en {case_data.get('id_caso', 'caso desconocido')} para {model}")
                    case_data[f'{model_key}_vn'] = 0
                
                # Precisión
                case_data[f'{model_key}_exactitud'] = self.extract_table_value(content, 'Exactitud', model)
                case_data[f'{model_key}_sensibilidad'] = self.extract_table_value(content, 'Sensibilidad', model)
                case_data[f'{model_key}_especificidad'] = self.extract_table_value(content, 'Especificidad', model)
                
                # Calidad - Claridad
                case_data[f'{model_key}_claridad_explicaciones'] = self.extract_table_value(content, 'Explicaciones Comprensibles', model)
                case_data[f'{model_key}_claridad_terminologia'] = self.extract_table_value(content, 'Uso Correcto de Terminología', model)
                case_data[f'{model_key}_claridad_estructura'] = self.extract_table_value(content, 'Estructura Lógica', model)
                
                # Calidad - Completitud
                case_data[f'{model_key}_completitud_lexica'] = self.extract_table_value(content, 'Cobertura de Similitud Léxica', model)
                case_data[f'{model_key}_completitud_estructural'] = self.extract_table_value(content, 'Cobertura de Similitud Estructural', model)
                case_data[f'{model_key}_completitud_funcional'] = self.extract_table_value(content, 'Cobertura de Similitud Funcional', model)
                case_data[f'{model_key}_completitud_estilo'] = self.extract_table_value(content, 'Cobertura de Similitud de Estilo', model)
                
                # Calidad - Coherencia
                case_data[f'{model_key}_coherencia_consistencia'] = self.extract_table_value(content, 'Consistencia entre Puntaje', model)
                case_data[f'{model_key}_coherencia_logica'] = self.extract_table_value(content, 'Lógica Interna', model)
            
            return case_data
                
        except Exception as e:
            print(f"Error procesando archivo {filepath}: {e}")
            return None
    
    def extract_value(self, content, pattern):
        """Extrae un valor usando regex"""
        match = re.search(pattern, content, re.IGNORECASE)
        return match.group(1) if match else None
    
    def extract_table_value(self, content, field_name, model_name):
        """
        Extrae un valor de una tabla para un modelo específico
        """
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            if field_name in line:
                # Buscar en las siguientes líneas los valores
                for j in range(i, min(i + 5, len(lines))):
                    current_line = lines[j]
                    
                    # Si la línea contiene todos los modelos
                    if all(m in current_line for m in ['Claude', 'GPT', 'Gemini', 'DeepSeek']):
                        # La siguiente línea debería tener los valores
                        if j + 1 < len(lines):
                            values_line = lines[j + 1]
                            values = re.findall(r'[\d.]+|N/A', values_line)
                            
                            # Mapear valores a modelos
                            model_index = {
                                'Claude Opus 4.1': 0,
                                'GPT-5': 1,
                                'Gemini 2.5 Pro': 2,
                                'DeepSeek-V3': 3
                            }
                            
                            if model_name in model_index and len(values) > model_index[model_name]:
                                value = values[model_index[model_name]]
                                if value != 'N/A':
                                    try:
                                        return float(value.replace('%', ''))
                                    except:
                                        return value
                                return None
                    
                    # Buscar patrón con tabs o espacios
                    if '\t' in current_line or '  ' in current_line:
                        parts = re.split(r'\t+|\s{2,}', current_line)
                        
                        if field_name in parts[0] if parts else '':
                            model_positions = {
                                'Claude Opus 4.1': 1,
                                'GPT-5': 2, 
                                'Gemini 2.5 Pro': 3,
                                'DeepSeek-V3': 4
                            }
                            
                            if model_name in model_positions and len(parts) > model_positions[model_name]:
                                value = parts[model_positions[model_name]].strip()
                                if value and value != 'N/A':
                                    try:
                                        return float(value.replace('%', ''))
                                    except:
                                        return value
        
        return None
    
    def load_all_cases(self):
        """Carga todos los archivos de casos"""
        print("Cargando archivos de casos...")
        
        files_processed = 0
        files_failed = 0
        java_count = 0
        csharp_count = 0
        
        if os.path.isdir(self.data_path):
            files = [f for f in os.listdir(self.data_path) if f.endswith(('.docx', '.txt', '.csv', '.md'))]
            total_files = len(files)
            
            print(f"Encontrados {total_files} archivos para procesar...")
            
            for filename in files:
                filepath = os.path.join(self.data_path, filename)
                print(f"Procesando: {filename}")
                
                case_data = self.parse_case_file(filepath)
                if case_data:
                    self.data.append(case_data)
                    files_processed += 1
                    
                    # Contar por lenguaje
                    if case_data.get('lenguaje') == 'Java':
                        java_count += 1
                    elif case_data.get('lenguaje') == 'C#':
                        csharp_count += 1
                else:
                    files_failed += 1
                    print(f"  ⚠️ No se pudo procesar: {filename}")
        
        # Crear DataFrame consolidado
        if self.data:
            self.df_consolidated = pd.DataFrame(self.data)
            
            # Llenar valores faltantes con valores por defecto
            for col in self.df_consolidated.columns:
                if 'exactitud' in col or 'sensibilidad' in col or 'especificidad' in col:
                    self.df_consolidated[col] = pd.to_numeric(self.df_consolidated[col], errors='coerce')
                elif any(metric in col for metric in ['vp', 'vn', 'fp', 'fn', 'coincidencia', 
                                                       'completitud', 'claridad', 'coherencia']):
                    self.df_consolidated[col] = pd.to_numeric(self.df_consolidated[col], errors='coerce').fillna(0)
                elif 'tiempo' in col or 'clasificacion' in col:
                    self.df_consolidated[col] = pd.to_numeric(self.df_consolidated[col], errors='coerce')
            
            print(f"\n✅ Procesamiento completado:")
            print(f"   - Archivos procesados exitosamente: {files_processed}")
            print(f"   - Archivos con errores: {files_failed}")
            print(f"   - Total de casos cargados: {len(self.data)}")
            print(f"   - Casos de Java: {java_count}")
            print(f"   - Casos de C#: {csharp_count}")
        else:
            print("❌ No se pudieron cargar datos de ningún archivo")
    
    def generate_general_comparison_matrix(self):
        """Genera la matriz de resultados generales por modelo"""
        print("\n=== MATRIZ DE RESULTADOS GENERALES POR MODELO ===")
        
        results = []
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            
            # Calcular precisión global
            precision_col = f'{model_key}_exactitud'
            if precision_col in self.df_consolidated.columns:
                precision = self.df_consolidated[precision_col].mean()
            else:
                precision = 0
            
            # Calcular tiempo promedio
            tiempo_col = f'{model_key}_tiempo'
            if tiempo_col in self.df_consolidated.columns:
                tiempo = self.df_consolidated[tiempo_col].mean()
            else:
                tiempo = 0
            
            # Calcular calidad global
            claridad_cols = [f'{model_key}_claridad_explicaciones', 
                           f'{model_key}_claridad_terminologia',
                           f'{model_key}_claridad_estructura']
            
            claridad = 0
            for col in claridad_cols:
                if col in self.df_consolidated.columns:
                    claridad += self.df_consolidated[col].mean() / len(claridad_cols)
            
            completitud_cols = [f'{model_key}_completitud_lexica',
                              f'{model_key}_completitud_estructural', 
                              f'{model_key}_completitud_funcional',
                              f'{model_key}_completitud_estilo']
            
            completitud = 0
            for col in completitud_cols:
                if col in self.df_consolidated.columns:
                    completitud += self.df_consolidated[col].mean()
            completitud = (completitud / 4) * 5  # Normalizar a escala 1-5
            
            coherencia_cols = [f'{model_key}_coherencia_consistencia',
                             f'{model_key}_coherencia_logica']
            
            coherencia = 0
            for col in coherencia_cols:
                if col in self.df_consolidated.columns:
                    coherencia += self.df_consolidated[col].mean() / len(coherencia_cols)
            
            calidad_global = claridad * 0.4 + completitud * 0.3 + coherencia * 0.3
            
            results.append({
                'Modelo': model,
                'Precisión Global (%)': round(precision, 2) if precision > 0 else 'N/A',
                'Tiempo Promedio (seg)': round(tiempo, 2) if tiempo > 0 else 'N/A',
                'Calidad Global (1-5)': round(calidad_global, 2) if calidad_global > 0 else 'N/A'
            })
        
        df_general = pd.DataFrame(results)
        
        # Calcular ranking
        numeric_cols = ['Precisión Global (%)', 'Tiempo Promedio (seg)', 'Calidad Global (1-5)']
        for col in numeric_cols:
            df_general[col] = pd.to_numeric(df_general[col], errors='coerce')
        
        df_general['Ranking Precisión'] = df_general['Precisión Global (%)'].rank(ascending=False, na_option='bottom')
        df_general['Ranking Tiempo'] = df_general['Tiempo Promedio (seg)'].rank(ascending=True, na_option='bottom')
        df_general['Ranking Calidad'] = df_general['Calidad Global (1-5)'].rank(ascending=False, na_option='bottom')
        df_general['Ranking General'] = (
            df_general['Ranking Precisión'] * 0.4 +
            df_general['Ranking Tiempo'] * 0.3 +
            df_general['Ranking Calidad'] * 0.3
        ).rank()
        
        print(df_general.to_string(index=False))
        return df_general
    
    def generate_precision_matrix(self):
        """Genera la matriz de precisión detallada"""
        print("\n=== MATRIZ DE PRECISIÓN DETALLADA ===")
        
        results = []
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            
            # Sumar totales de matriz de confusión TAL COMO ESTÁN EN LOS DATOS
            vp_col = f'{model_key}_vp'
            vn_col = f'{model_key}_vn'
            fp_col = f'{model_key}_fp'
            fn_col = f'{model_key}_fn'
            
            vp_total = self.df_consolidated[vp_col].sum() if vp_col in self.df_consolidated.columns else 0
            vn_total = self.df_consolidated[vn_col].sum() if vn_col in self.df_consolidated.columns else 0
            fp_total = self.df_consolidated[fp_col].sum() if fp_col in self.df_consolidated.columns else 0
            fn_total = self.df_consolidated[fn_col].sum() if fn_col in self.df_consolidated.columns else 0
            
            # Calcular métricas usando los valores TAL COMO ESTÁN
            total = vp_total + vn_total + fp_total + fn_total
            
            if total > 0:
                exactitud = ((vp_total + vn_total) / total) * 100
            else:
                exactitud = 0
            
            if (vp_total + fn_total) > 0:
                sensibilidad = (vp_total / (vp_total + fn_total)) * 100
            else:
                sensibilidad = 0
            
            if (vn_total + fp_total) > 0:
                especificidad = (vn_total / (vn_total + fp_total)) * 100
            else:
                especificidad = 'N/A'
            
            results.append({
                'Modelo': model,
                'Exactitud (%)': round(exactitud, 2),
                'Sensibilidad (%)': round(sensibilidad, 2),
                'Especificidad (%)': round(especificidad, 2) if especificidad != 'N/A' else 'N/A',
                'VP': int(vp_total),
                'VN': int(vn_total),
                'FP': int(fp_total),
                'FN': int(fn_total)
            })
        
        df_precision = pd.DataFrame(results)
        print(df_precision.to_string(index=False))
        return df_precision
    
    def generate_similarity_category_matrix(self):
        """Genera matriz por categoría de similitud"""
        print("\n=== MATRIZ POR CATEGORÍA DE SIMILITUD ===")
        
        categories = self.df_consolidated['categoria_similitud'].unique() if 'categoria_similitud' in self.df_consolidated.columns else []
        results = []
        
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            row = {'Modelo': model}
            
            for category in categories:
                if category:
                    category_data = self.df_consolidated[self.df_consolidated['categoria_similitud'] == category]
                    coincidencia_col = f'{model_key}_coincidencia'
                    
                    if len(category_data) > 0 and coincidencia_col in self.df_consolidated.columns:
                        accuracy = category_data[coincidencia_col].mean() * 100
                        row[f'{category} (%)'] = round(accuracy, 2)
                    else:
                        row[f'{category} (%)'] = 'N/A'
            
            results.append(row)
        
        df_category = pd.DataFrame(results)
        print(df_category.to_string(index=False))
        return df_category
    
    def generate_language_matrix(self):
        """Genera matriz por lenguaje de programación (Java y C#)"""
        print("\n=== MATRIZ POR LENGUAJE DE PROGRAMACIÓN ===")
        
        results = []
        
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            
            # Java
            java_data = self.df_consolidated[self.df_consolidated['lenguaje'] == 'Java'] if 'lenguaje' in self.df_consolidated.columns else pd.DataFrame()
            exactitud_col = f'{model_key}_exactitud'
            tiempo_col = f'{model_key}_tiempo'
            
            java_precision = java_data[exactitud_col].mean() if len(java_data) > 0 and exactitud_col in java_data.columns else 0
            java_tiempo = java_data[tiempo_col].mean() if len(java_data) > 0 and tiempo_col in java_data.columns else 0
            
            # C#
            csharp_data = self.df_consolidated[self.df_consolidated['lenguaje'] == 'C#'] if 'lenguaje' in self.df_consolidated.columns else pd.DataFrame()
            csharp_precision = csharp_data[exactitud_col].mean() if len(csharp_data) > 0 and exactitud_col in csharp_data.columns else 0
            csharp_tiempo = csharp_data[tiempo_col].mean() if len(csharp_data) > 0 and tiempo_col in csharp_data.columns else 0
            
            results.append({
                'Modelo': model,
                'Java - Precisión (%)': round(java_precision, 2) if java_precision > 0 else 'N/A',
                'Java - Tiempo (seg)': round(java_tiempo, 2) if java_tiempo > 0 else 'N/A',
                'C# - Precisión (%)': round(csharp_precision, 2) if csharp_precision > 0 else 'N/A',
                'C# - Tiempo (seg)': round(csharp_tiempo, 2) if csharp_tiempo > 0 else 'N/A'
            })
        
        df_language = pd.DataFrame(results)
        print(df_language.to_string(index=False))
        return df_language
    
    def generate_time_response_matrix(self):
        """Genera matriz de tiempos de respuesta"""
        print("\n=== MATRIZ DE TIEMPOS DE RESPUESTA ===")
        
        results = []
        
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            tiempo_col = f'{model_key}_tiempo'
            
            if tiempo_col in self.df_consolidated.columns:
                tiempos = self.df_consolidated[tiempo_col].dropna()
                
                if len(tiempos) > 0:
                    results.append({
                        'Modelo': model,
                        'Tiempo Promedio (seg)': round(tiempos.mean(), 2),
                        'Desviación Estándar': round(tiempos.std(), 3),
                        'Coef. Variación (%)': round((tiempos.std() / tiempos.mean()) * 100, 2) if tiempos.mean() > 0 else 'N/A',
                        'Tiempo Mín': round(tiempos.min(), 2),
                        'Tiempo Máx': round(tiempos.max(), 2)
                    })
                else:
                    results.append({
                        'Modelo': model,
                        'Tiempo Promedio (seg)': 'N/A',
                        'Desviación Estándar': 'N/A',
                        'Coef. Variación (%)': 'N/A',
                        'Tiempo Mín': 'N/A',
                        'Tiempo Máx': 'N/A'
                    })
        
        df_time = pd.DataFrame(results)
        print(df_time.to_string(index=False))
        return df_time
    
    def generate_quality_matrix(self):
        """Genera matriz de calidad de justificaciones"""
        print("\n=== MATRIZ DE CALIDAD DE JUSTIFICACIONES ===")
        
        results = []
        
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            
            # Claridad (promedio de 3 indicadores)
            claridad_cols = [f'{model_key}_claridad_explicaciones',
                           f'{model_key}_claridad_terminologia',
                           f'{model_key}_claridad_estructura']
            
            claridad = 0
            count = 0
            for col in claridad_cols:
                if col in self.df_consolidated.columns:
                    claridad += self.df_consolidated[col].mean()
                    count += 1
            claridad = claridad / count if count > 0 else 0
            
            # Completitud (suma de 4 indicadores, máximo 4)
            completitud_cols = [f'{model_key}_completitud_lexica',
                              f'{model_key}_completitud_estructural',
                              f'{model_key}_completitud_funcional',
                              f'{model_key}_completitud_estilo']
            
            completitud = 0
            for col in completitud_cols:
                if col in self.df_consolidated.columns:
                    completitud += self.df_consolidated[col].mean()
            
            # Coherencia (promedio de 2 indicadores)
            coherencia_cols = [f'{model_key}_coherencia_consistencia',
                             f'{model_key}_coherencia_logica']
            
            coherencia = 0
            count = 0
            for col in coherencia_cols:
                if col in self.df_consolidated.columns:
                    coherencia += self.df_consolidated[col].mean()
                    count += 1
            coherencia = coherencia / count if count > 0 else 0
            
            # Puntuación total
            puntuacion_total = claridad + (completitud * 5/4) + coherencia
            
            results.append({
                'Modelo': model,
                'Claridad (1-5)': round(claridad, 2) if claridad > 0 else 'N/A',
                'Completitud (0-4)': round(completitud, 2) if completitud > 0 else 'N/A',
                'Coherencia (1-5)': round(coherencia, 2) if coherencia > 0 else 'N/A',
                'Puntuación Total': round(puntuacion_total, 2) if puntuacion_total > 0 else 'N/A'
            })
        
        df_quality = pd.DataFrame(results)
        print(df_quality.to_string(index=False))
        return df_quality
    
    def perform_statistical_analysis(self):
        """Realiza análisis estadístico comparativo"""
        print("\n=== ANÁLISIS ESTADÍSTICO COMPARATIVO ===")
        
        # Preparar datos para ANOVA
        precision_data = []
        time_data = []
        quality_data = []
        
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            
            # Precisión
            exactitud_col = f'{model_key}_exactitud'
            if exactitud_col in self.df_consolidated.columns:
                data = self.df_consolidated[exactitud_col].dropna().values
                if len(data) > 0:
                    precision_data.append(data)
            
            # Tiempo
            tiempo_col = f'{model_key}_tiempo'
            if tiempo_col in self.df_consolidated.columns:
                data = self.df_consolidated[tiempo_col].dropna().values
                if len(data) > 0:
                    time_data.append(data)
            
            # Calidad combinada
            quality_scores = []
            for idx in range(len(self.df_consolidated)):
                score = 0
                count = 0
                for col_type in ['claridad_explicaciones', 'claridad_terminologia', 'claridad_estructura']:
                    col = f'{model_key}_{col_type}'
                    if col in self.df_consolidated.columns:
                        val = self.df_consolidated.iloc[idx][col]
                        if pd.notna(val):
                            score += val
                            count += 1
                if count > 0:
                    quality_scores.append(score / count)
            if quality_scores:
                quality_data.append(quality_scores)
        
        results = []
        
        # ANOVA para precisión
        if len(precision_data) >= 2:
            f_stat_precision, p_val_precision = stats.f_oneway(*precision_data)
            results.append({
                'Prueba Estadística': 'ANOVA - Precisión',
                'Valor Estadístico': round(f_stat_precision, 4),
                'p-valor': round(p_val_precision, 6),
                'Significancia': 'Sí' if p_val_precision < 0.05 else 'No',
                'Conclusión': 'Diferencias significativas' if p_val_precision < 0.05 else 'Sin diferencias significativas'
            })
        
        # ANOVA para tiempo
        if len(time_data) >= 2:
            f_stat_time, p_val_time = stats.f_oneway(*time_data)
            results.append({
                'Prueba Estadística': 'ANOVA - Tiempo',
                'Valor Estadístico': round(f_stat_time, 4),
                'p-valor': round(p_val_time, 6),
                'Significancia': 'Sí' if p_val_time < 0.05 else 'No',
                'Conclusión': 'Diferencias significativas' if p_val_time < 0.05 else 'Sin diferencias significativas'
            })
        
        # Kruskal-Wallis para calidad (ordinal)
        if len(quality_data) >= 2:
            h_stat_quality, p_val_quality = stats.kruskal(*quality_data)
            results.append({
                'Prueba Estadística': 'Kruskal-Wallis - Calidad',
                'Valor Estadístico': round(h_stat_quality, 4),
                'p-valor': round(p_val_quality, 6),
                'Significancia': 'Sí' if p_val_quality < 0.05 else 'No',
                'Conclusión': 'Diferencias significativas' if p_val_quality < 0.05 else 'Sin diferencias significativas'
            })
        
        if results:
            df_stats = pd.DataFrame(results)
            print(df_stats.to_string(index=False))
            
            # Si hay diferencias significativas, realizar post-hoc
            if any(r['Significancia'] == 'Sí' for r in results):
                print("\n=== ANÁLISIS POST-HOC (Tukey HSD) ===")
                if len(precision_data) >= 2 and len(time_data) >= 2:
                    self.perform_posthoc_analysis(precision_data, time_data)
            
            return df_stats
        else:
            print("No hay suficientes datos para realizar análisis estadístico")
            return pd.DataFrame()
    
    def perform_posthoc_analysis(self, precision_data, time_data):
        """Realiza análisis post-hoc si es necesario"""
        try:
            from scipy.stats import tukey_hsd
            
            # Tukey HSD para precisión
            if len(precision_data) >= 2:
                print("\nPost-hoc para Precisión:")
                res_precision = tukey_hsd(*precision_data)
                print(res_precision)
            
            # Tukey HSD para tiempo
            if len(time_data) >= 2:
                print("\nPost-hoc para Tiempo:")
                res_time = tukey_hsd(*time_data)
                print(res_time)
        except Exception as e:
            print(f"No se pudo realizar análisis post-hoc: {e}")
    
    def generate_visualizations(self):
        """Genera visualizaciones de los resultados en archivos separados"""
        print("\n=== GENERANDO VISUALIZACIONES ===")
        
        # Configurar estilo
        sns.set_style("whitegrid")
        
        try:
            # 1. GRÁFICO DE PRECISIÓN GLOBAL
            fig1, ax1 = plt.subplots(figsize=(10, 6))
            precision_means = []
            model_names = []
            for model in self.models:
                model_key = model.replace(' ', '_').replace('.', '_')
                exactitud_col = f'{model_key}_exactitud'
                if exactitud_col in self.df_consolidated.columns:
                    mean_val = self.df_consolidated[exactitud_col].mean()
                    if pd.notna(mean_val):
                        precision_means.append(mean_val)
                        model_names.append(model)
            
            if precision_means:
                bars = ax1.bar(model_names, precision_means, color=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728'])
                ax1.set_title('Precisión Global por Modelo de IA', fontsize=14, fontweight='bold')
                ax1.set_ylabel('Precisión (%)', fontsize=12)
                ax1.set_xlabel('Modelo', fontsize=12)
                ax1.set_ylim(0, 110)
                for i, (bar, v) in enumerate(zip(bars, precision_means)):
                    ax1.text(bar.get_x() + bar.get_width()/2, v + 1, f'{v:.1f}%', 
                           ha='center', va='bottom', fontweight='bold')
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                plt.savefig('01_precision_global.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # 2. BOXPLOT DE TIEMPOS DE RESPUESTA
            fig2, ax2 = plt.subplots(figsize=(10, 6))
            time_data = []
            time_labels = []
            for model in self.models:
                model_key = model.replace(' ', '_').replace('.', '_')
                tiempo_col = f'{model_key}_tiempo'
                if tiempo_col in self.df_consolidated.columns:
                    data = self.df_consolidated[tiempo_col].dropna().values
                    if len(data) > 0:
                        time_data.append(data)
                        time_labels.append(model)
            
            if time_data:
                bp = ax2.boxplot(time_data, labels=time_labels, patch_artist=True)
                colors = ['#FFB6C1', '#87CEEB', '#98FB98', '#DDA0DD']
                for patch, color in zip(bp['boxes'], colors):
                    patch.set_facecolor(color)
                ax2.set_title('Distribución de Tiempos de Respuesta', fontsize=14, fontweight='bold')
                ax2.set_ylabel('Tiempo (segundos)', fontsize=12)
                ax2.set_xlabel('Modelo', fontsize=12)
                ax2.grid(True, alpha=0.3)
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                plt.savefig('02_tiempos_respuesta.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # 3. CALIDAD DE JUSTIFICACIONES
            fig3, ax3 = plt.subplots(figsize=(10, 6))
            quality_scores = []
            quality_labels = []
            for model in self.models:
                model_key = model.replace(' ', '_').replace('.', '_')
                claridad_col = f'{model_key}_claridad_explicaciones'
                if claridad_col in self.df_consolidated.columns:
                    mean_val = self.df_consolidated[claridad_col].mean()
                    if pd.notna(mean_val):
                        quality_scores.append(mean_val)
                        quality_labels.append(model)
            
            if quality_scores:
                bars = ax3.bar(quality_labels, quality_scores, color=['#2E8B57', '#4682B4', '#DC143C', '#FF8C00'])
                ax3.set_title('Calidad de Justificaciones por Modelo', fontsize=14, fontweight='bold')
                ax3.set_ylabel('Puntuación (1-5)', fontsize=12)
                ax3.set_xlabel('Modelo', fontsize=12)
                ax3.set_ylim(0, 5.5)
                ax3.axhline(y=3, color='gray', linestyle='--', alpha=0.5, label='Promedio esperado')
                for bar, v in zip(bars, quality_scores):
                    ax3.text(bar.get_x() + bar.get_width()/2, v + 0.05, f'{v:.2f}', 
                           ha='center', va='bottom', fontweight='bold')
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                plt.savefig('03_calidad_justificaciones.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            # 4. COMPARACIÓN POR LENGUAJE (JAVA vs C#)
            if 'lenguaje' in self.df_consolidated.columns:
                fig4, ax4 = plt.subplots(figsize=(12, 6))
                java_precision = []
                csharp_precision = []
                lang_models = []
                
                for model in self.models:
                    model_key = model.replace(' ', '_').replace('.', '_')
                    exactitud_col = f'{model_key}_exactitud'
                    
                    if exactitud_col in self.df_consolidated.columns:
                        java_data = self.df_consolidated[self.df_consolidated['lenguaje'] == 'Java']
                        csharp_data = self.df_consolidated[self.df_consolidated['lenguaje'] == 'C#']
                        
                        java_mean = java_data[exactitud_col].mean() if len(java_data) > 0 else 0
                        csharp_mean = csharp_data[exactitud_col].mean() if len(csharp_data) > 0 else 0
                        
                        if pd.notna(java_mean) and pd.notna(csharp_mean):
                            java_precision.append(java_mean)
                            csharp_precision.append(csharp_mean)
                            lang_models.append(model)
                
                if java_precision and csharp_precision:
                    x = np.arange(len(lang_models))
                    width = 0.35
                    bars1 = ax4.bar(x - width/2, java_precision, width, label='Java', color='#FF6B6B')
                    bars2 = ax4.bar(x + width/2, csharp_precision, width, label='C#', color='#4ECDC4')
                    
                    ax4.set_xlabel('Modelo', fontsize=12)
                    ax4.set_ylabel('Precisión (%)', fontsize=12)
                    ax4.set_title('Precisión por Lenguaje de Programación', fontsize=14, fontweight='bold')
                    ax4.set_xticks(x)
                    ax4.set_xticklabels(lang_models, rotation=45, ha='right')
                    ax4.legend()
                    
                    # Agregar valores en las barras
                    for bars in [bars1, bars2]:
                        for bar in bars:
                            height = bar.get_height()
                            ax4.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                                   f'{height:.1f}', ha='center', va='bottom', fontsize=9)
                    
                    plt.tight_layout()
                    plt.savefig('04_precision_por_lenguaje.png', dpi=300, bbox_inches='tight')
                    plt.close()
            
            # 5. MATRIZ DE CONFUSIÓN CONSOLIDADA
            fig5, ax5 = plt.subplots(figsize=(10, 6))
            confusion_matrix = np.zeros((len(self.models), 4))
            for i, model in enumerate(self.models):
                model_key = model.replace(' ', '_').replace('.', '_')
                confusion_matrix[i, 0] = self.df_consolidated[f'{model_key}_vp'].sum() if f'{model_key}_vp' in self.df_consolidated.columns else 0
                confusion_matrix[i, 1] = self.df_consolidated[f'{model_key}_vn'].sum() if f'{model_key}_vn' in self.df_consolidated.columns else 0
                confusion_matrix[i, 2] = self.df_consolidated[f'{model_key}_fp'].sum() if f'{model_key}_fp' in self.df_consolidated.columns else 0
                confusion_matrix[i, 3] = self.df_consolidated[f'{model_key}_fn'].sum() if f'{model_key}_fn' in self.df_consolidated.columns else 0
            
            im = ax5.imshow(confusion_matrix, cmap='YlOrRd', aspect='auto')
            ax5.set_xticks(range(4))
            ax5.set_xticklabels(['VP', 'VN', 'FP', 'FN'])
            ax5.set_yticks(range(len(self.models)))
            ax5.set_yticklabels(self.models)
            ax5.set_title('Matriz de Confusión Consolidada', fontsize=14, fontweight='bold')
            
            # Agregar valores en el heatmap
            for i in range(len(self.models)):
                for j in range(4):
                    text = ax5.text(j, i, f'{int(confusion_matrix[i, j])}',
                                   ha="center", va="center", color="black", fontweight='bold')
            
            plt.colorbar(im, ax=ax5)
            plt.tight_layout()
            plt.savefig('05_matriz_confusion.png', dpi=300, bbox_inches='tight')
            plt.close()
            
            # 6. HEATMAP DE RENDIMIENTO POR CATEGORÍA
            if 'categoria_similitud' in self.df_consolidated.columns:
                fig6, ax6 = plt.subplots(figsize=(10, 6))
                categories = self.df_consolidated['categoria_similitud'].dropna().unique()
                if len(categories) > 0:
                    category_performance = []
                    for model in self.models:
                        model_key = model.replace(' ', '_').replace('.', '_')
                        coincidencia_col = f'{model_key}_coincidencia'
                        model_perf = []
                        
                        for cat in categories:
                            cat_data = self.df_consolidated[self.df_consolidated['categoria_similitud'] == cat]
                            if len(cat_data) > 0 and coincidencia_col in cat_data.columns:
                                mean_val = cat_data[coincidencia_col].mean() * 100
                                model_perf.append(mean_val if pd.notna(mean_val) else 0)
                            else:
                                model_perf.append(0)
                        category_performance.append(model_perf)
                    
                    if category_performance:
                        im2 = ax6.imshow(category_performance, cmap='RdYlGn', aspect='auto', vmin=0, vmax=100)
                        ax6.set_xticks(range(len(categories)))
                        ax6.set_xticklabels(categories, rotation=45, ha='right')
                        ax6.set_yticks(range(len(self.models)))
                        ax6.set_yticklabels(self.models)
                        ax6.set_title('Rendimiento por Categoría de Similitud (%)', fontsize=14, fontweight='bold')
                        
                        # Agregar valores en el heatmap
                        for i in range(len(self.models)):
                            for j in range(len(categories)):
                                ax6.text(j, i, f'{category_performance[i][j]:.1f}',
                                       ha='center', va='center', color='black', fontweight='bold')
                        
                        plt.colorbar(im2, ax=ax6, label='Precisión (%)')
                        plt.tight_layout()
                        plt.savefig('06_rendimiento_por_categoria.png', dpi=300, bbox_inches='tight')
                        plt.close()
            
            # 7. GRÁFICO COMPARATIVO DE TIEMPOS POR LENGUAJE
            if 'lenguaje' in self.df_consolidated.columns:
                fig7, (ax7a, ax7b) = plt.subplots(1, 2, figsize=(14, 6))
                
                # Tiempos para Java
                java_times = []
                # Tiempos para C#
                csharp_times = []
                model_labels = []
                
                for model in self.models:
                    model_key = model.replace(' ', '_').replace('.', '_')
                    tiempo_col = f'{model_key}_tiempo'
                    
                    if tiempo_col in self.df_consolidated.columns:
                        java_data = self.df_consolidated[self.df_consolidated['lenguaje'] == 'Java']
                        csharp_data = self.df_consolidated[self.df_consolidated['lenguaje'] == 'C#']
                        
                        java_time = java_data[tiempo_col].mean() if len(java_data) > 0 else 0
                        csharp_time = csharp_data[tiempo_col].mean() if len(csharp_data) > 0 else 0
                        
                        if pd.notna(java_time) and pd.notna(csharp_time):
                            java_times.append(java_time)
                            csharp_times.append(csharp_time)
                            model_labels.append(model)
                
                # Gráfico de barras para tiempos Java
                if java_times:
                    bars = ax7a.bar(model_labels, java_times, color='#FF6B6B')
                    ax7a.set_title('Tiempo Promedio - Java', fontsize=12, fontweight='bold')
                    ax7a.set_ylabel('Tiempo (segundos)', fontsize=10)
                    ax7a.set_xticklabels(model_labels, rotation=45, ha='right', fontsize=9)
                    for bar, v in zip(bars, java_times):
                        ax7a.text(bar.get_x() + bar.get_width()/2, v + 0.2, f'{v:.1f}s',
                                ha='center', va='bottom', fontsize=9)
                
                # Gráfico de barras para tiempos C#
                if csharp_times:
                    bars = ax7b.bar(model_labels, csharp_times, color='#4ECDC4')
                    ax7b.set_title('Tiempo Promedio - C#', fontsize=12, fontweight='bold')
                    ax7b.set_ylabel('Tiempo (segundos)', fontsize=10)
                    ax7b.set_xticklabels(model_labels, rotation=45, ha='right', fontsize=9)
                    for bar, v in zip(bars, csharp_times):
                        ax7b.text(bar.get_x() + bar.get_width()/2, v + 0.2, f'{v:.1f}s',
                                ha='center', va='bottom', fontsize=9)
                
                plt.suptitle('Comparación de Tiempos de Respuesta por Lenguaje', fontsize=14, fontweight='bold')
                plt.tight_layout()
                plt.savefig('07_tiempos_por_lenguaje.png', dpi=300, bbox_inches='tight')
                plt.close()
            
            print("✅ Visualizaciones guardadas:")
            print("   📊 01_precision_global.png")
            print("   📊 02_tiempos_respuesta.png")
            print("   📊 03_calidad_justificaciones.png")
            print("   📊 04_precision_por_lenguaje.png")
            print("   📊 05_matriz_confusion.png")
            print("   📊 06_rendimiento_por_categoria.png")
            print("   📊 07_tiempos_por_lenguaje.png")
            
        except Exception as e:
            print(f"⚠️ Error generando visualizaciones: {e}")
            import traceback
            traceback.print_exc()
    
    def export_results(self):
        """Exporta todos los resultados a Excel"""
        print("\n=== EXPORTANDO RESULTADOS A EXCEL ===")
        
        try:
            with pd.ExcelWriter('resultados_comparacion_modelos.xlsx', engine='openpyxl') as writer:
                # Datos consolidados
                self.df_consolidated.to_excel(writer, sheet_name='Datos_Completos', index=False)
                
                # Matrices de comparación
                self.generate_general_comparison_matrix().to_excel(writer, sheet_name='Comparacion_General', index=False)
                self.generate_precision_matrix().to_excel(writer, sheet_name='Precision_Detallada', index=False)
                self.generate_similarity_category_matrix().to_excel(writer, sheet_name='Por_Categoria', index=False)
                self.generate_language_matrix().to_excel(writer, sheet_name='Por_Lenguaje', index=False)
                self.generate_time_response_matrix().to_excel(writer, sheet_name='Tiempos_Respuesta', index=False)
                self.generate_quality_matrix().to_excel(writer, sheet_name='Calidad_Justificaciones', index=False)
                
                # Análisis estadístico
                stats_df = self.perform_statistical_analysis()
                if not stats_df.empty:
                    stats_df.to_excel(writer, sheet_name='Analisis_Estadistico', index=False)
                
            print("✅ Resultados exportados a 'resultados_comparacion_modelos.xlsx'")
        except Exception as e:
            print(f"⚠️ Error exportando resultados: {e}")
    
    def generate_summary_report(self):
        """Genera un reporte resumido con las conclusiones principales"""
        print("\n" + "="*80)
        print("REPORTE EJECUTIVO - COMPARACIÓN DE MODELOS DE IA")
        print("="*80)
        
        # Mejor modelo por categoría
        print("\n### MEJORES MODELOS POR CATEGORÍA ###")
        
        # Precisión
        best_precision_model = None
        best_precision = 0
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            exactitud_col = f'{model_key}_exactitud'
            if exactitud_col in self.df_consolidated.columns:
                precision = self.df_consolidated[exactitud_col].mean()
                if pd.notna(precision) and precision > best_precision:
                    best_precision = precision
                    best_precision_model = model
        
        if best_precision_model:
            print(f"✓ Mejor Precisión: {best_precision_model} ({best_precision:.2f}%)")
        
        # Tiempo
        best_time_model = None
        best_time = float('inf')
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            tiempo_col = f'{model_key}_tiempo'
            if tiempo_col in self.df_consolidated.columns:
                time = self.df_consolidated[tiempo_col].mean()
                if pd.notna(time) and time < best_time:
                    best_time = time
                    best_time_model = model
        
        if best_time_model:
            print(f"✓ Mejor Tiempo de Respuesta: {best_time_model} ({best_time:.2f} seg)")
        
        # Calidad
        best_quality_model = None
        best_quality = 0
        for model in self.models:
            model_key = model.replace(' ', '_').replace('.', '_')
            claridad_col = f'{model_key}_claridad_explicaciones'
            if claridad_col in self.df_consolidated.columns:
                quality = self.df_consolidated[claridad_col].mean()
                if pd.notna(quality) and quality > best_quality:
                    best_quality = quality
                    best_quality_model = model
        
        if best_quality_model:
            print(f"✓ Mejor Calidad de Justificaciones: {best_quality_model} ({best_quality:.2f}/5)")
        
        # Estadísticas generales
        print("\n### ESTADÍSTICAS GENERALES ###")
        print(f"• Total de casos evaluados: {len(self.df_consolidated)}")
        
        if 'lenguaje' in self.df_consolidated.columns:
            java_count = len(self.df_consolidated[self.df_consolidated['lenguaje'] == 'Java'])
            csharp_count = len(self.df_consolidated[self.df_consolidated['lenguaje'] == 'C#'])
            print(f"• Casos de Java: {java_count}")
            print(f"• Casos de C#: {csharp_count}")
        
        # Distribución por categoría
        if 'categoria_similitud' in self.df_consolidated.columns:
            print("\n### DISTRIBUCIÓN POR CATEGORÍA DE SIMILITUD ###")
            categories = self.df_consolidated['categoria_similitud'].dropna().unique()
            for category in categories:
                count = len(self.df_consolidated[self.df_consolidated['categoria_similitud'] == category])
                percentage = (count / len(self.df_consolidated)) * 100
                print(f"• {category}: {count} casos ({percentage:.1f}%)")
        
        print("\n" + "="*80)

def main():
    """Función principal para ejecutar el análisis"""
    
    # Configurar rutas
    print("="*80)
    print("SISTEMA DE ANÁLISIS COMPARATIVO DE MODELOS DE IA")
    print("Detección de Similitudes en Código Fuente")
    print("="*80)
    
    # Verificar instalación de python-docx
    try:
        from docx import Document
    except ImportError:
        print("\n⚠️ La biblioteca 'python-docx' no está instalada.")
        print("Por favor, instálela ejecutando:")
        print("   python -m pip install python-docx")
        return
    
    # Solicitar ruta de los archivos
    print("\nEste script procesará archivos con el patrón:")
    print("  - CASE_CODE_JAVA_XXX")
    print("  - CASE_CODE_C#_XXX")
    
    data_path = input("\nIngrese la ruta de la carpeta con los 120 casos .docx (o presione Enter para usar 'casos/'): ")
    if not data_path:
        data_path = "casos/"
    
    # Verificar que la carpeta existe
    if not os.path.exists(data_path):
        print(f"\n❌ La carpeta '{data_path}' no existe.")
        print("Por favor, cree la carpeta y coloque los archivos .docx dentro.")
        return
    
    # Crear instancia del analizador
    analyzer = ModelComparison(data_path)
    
    try:
        # Cargar todos los casos
        analyzer.load_all_cases()
        
        if len(analyzer.data) == 0:
            print("\n❌ No se pudieron cargar datos. Verifique que los archivos estén en el formato correcto.")
            print("Los archivos deben tener nombres como:")
            print("  - CASE_CODE_JAVA_001.docx")
            print("  - CASE_CODE_C#_001.docx")
            return
        
        # Generar todas las matrices
        print("\n" + "="*50)
        print("GENERANDO MATRICES DE COMPARACIÓN")
        print("="*50)
        
        analyzer.generate_general_comparison_matrix()
        analyzer.generate_precision_matrix()
        analyzer.generate_similarity_category_matrix()
        analyzer.generate_language_matrix()
        analyzer.generate_time_response_matrix()
        analyzer.generate_quality_matrix()
        
        # Análisis estadístico
        analyzer.perform_statistical_analysis()
        
        # Generar visualizaciones
        analyzer.generate_visualizations()
        
        # Exportar resultados
        analyzer.export_results()
        
        # Generar reporte resumido
        analyzer.generate_summary_report()
        
        print("\n✅ ANÁLISIS COMPLETADO EXITOSAMENTE")
        print("\n📁 ARCHIVOS GENERADOS:")
        print("\n📊 Excel con matrices:")
        print("   • resultados_comparacion_modelos.xlsx")
        print("\n📈 Gráficos individuales:")
        print("   • 01_precision_global.png - Comparación de precisión")
        print("   • 02_tiempos_respuesta.png - Distribución de tiempos")
        print("   • 03_calidad_justificaciones.png - Puntuación de calidad")
        print("   • 04_precision_por_lenguaje.png - Java vs C#")
        print("   • 05_matriz_confusion.png - Matriz de confusión")
        print("   • 06_rendimiento_por_categoria.png - Heatmap de categorías")
        print("   • 07_tiempos_por_lenguaje.png - Tiempos Java vs C#")
        
        print("\n💡 RECOMENDACIONES:")
        print("1. Revise el archivo Excel para análisis detallado de datos")
        print("2. Use los gráficos individuales para su presentación de tesis")
        print("3. El análisis estadístico está en la hoja 'Analisis_Estadistico'")
        print("4. Los p-valores indican si hay diferencias significativas entre modelos")
        
    except Exception as e:
        print(f"\n❌ Error durante el análisis: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()